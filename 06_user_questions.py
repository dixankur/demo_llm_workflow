import os
import time
import openai
import pandas as pd
from dotenv import load_dotenv, find_dotenv

from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter

import pinecone
from langchain.vectorstores import Pinecone
from langchain.embeddings.openai import OpenAIEmbeddings

from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA

import docx
from docx.shared import Pt

# Read environment variables from .env file
load_dotenv(find_dotenv(), override=True)

# Set the Open AI key from environment variables
openai.api_key = os.getenv("OPENAI_API_KEY")

# Load the transcript data
loader = PyPDFLoader("./data/first-call-SNOW-ticket-categorization.pdf")
pages = loader.load()

# Split the document into chunks
data_splitter = RecursiveCharacterTextSplitter(
    chunk_size=512, 
    chunk_overlap=64, 
    length_function = len,
    add_start_index = True,
)
chunks = data_splitter.split_documents(pages)

print(f'Total number of chunks generated: {len(chunks)}')

# Initialize open ai embedings
oi_embeddings = OpenAIEmbeddings()

# Initialize pinecone
index_name = 'transcript-index'
pinecone.init(api_key=os.getenv('PINECONE_API_KEY'), environment=os.getenv('PINECONE_ENV'))

# check if pinecone index already exists
if index_name in pinecone.list_indexes():
    print(f'Index {index_name} already exists. Loading embeddings...')
    vector_store = Pinecone.from_existing_index(index_name, oi_embeddings)
    print('Ok')
else:
    print(f'Creating index {index_name} and embeddings ...')
    pinecone.create_index(index_name, dimension=1536, metric='cosine')
    time.sleep(120)
    vector_store = Pinecone.from_documents(chunks, oi_embeddings, index_name=index_name)
    print('Ok')

llm = ChatOpenAI(model='gpt-3.5-turbo', temperature=0.7)
retriever = vector_store.as_retriever(search_type='similarity', search_kwargs={'k':7})

chain = RetrievalQA.from_chain_type(llm=llm, chain_type='stuff', retriever=retriever)

question_list = [
    (1, 'How many Speaker attended the meeting?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'The various participants or speakers in the transcript are in format like "Speaker n (mm:ss):" '
      'where n is the speaker number and mm is time in minutes into the meeting and ss is seconds into the meeting. '
      'Based on the different number of speakers in the transcript, calculate approximately how many speakers may have attended the meeting? '
      'You must provide us an answer with your best possible guess.'
     )),
     (2, 'How would you describe the use case in one sentence?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'How would you describe the use case in one sentence?'
     )),
     (3, 'Describe the name of the use case in 3-5 words?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Describe the name of the use case in 3-5 words?'
     )),
     (4, 'What is the target feature?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, please tell what is the target feature?'
     )),
     (5, 'What is the unit of analysis?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, please tell What is the unit of analysis?'
     )),
     (6, 'What is the current process?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, please tell what is the current process being used? '
      'Please tell about current process only, do not explain about use of machine learning and AI in future.'
     )),
     (7, 'Where is the data located?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, can you provide some insights about where the data might be located and in which format?'
     )),
    (8, 'What is the desired prediction frequency?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, can you tell what is the desired prediction frequency of the proposed machine learning or AI model?'
     )),
     (9, 'Did anyone mention any business constraints?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, can you tell about any business constraints being discussed?'
     )),
     (10, 'How long should the prediction be retained in the data warehouse?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, can you tell the proposed time for the predictions from the model to be retained?'
     )),
     (11, 'What is the success criteria?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, can you tell what success criterias are being discussed for the provided problem or use case?'
     )),
     (12, 'What features should be included in the dataset?', 
     ('You are an experienced data scientist and experienced machine learning and artificial intelligence expert. '
      'You are provided with a transcript of a meeting where a use case is being discussed about possible use of '
      'machine learning and artificial intelligence to solve a user problem. '
      'Based on above context, can you tell which features are proposed to be included in the dataset for the input '
      'to the machine learning model?'
     )),
]

total_text = ''
for seq, question, prompt in question_list:
    print('-'*30)
    print(f"Question {seq}: {question}")
    total_text = total_text + f"Question {seq}: {question} \n" 
    print("  ")
    print(f"Prompt used: {prompt}")
    answer = chain.run(prompt)
    print("  ")
    print(f"Answer: {answer}\n")
    total_text = total_text + f"Answer: {answer} \n"

print("======================================================")
print(total_text)

use_case_prompt = f'''
    Based on the set of question and answers provided below create a use case document.
    The use case document should include at least the following sections:
    1. Introduction
    2. Use case description
    3. Current process
    4. Busines constraints
    5. Success criteria
    6. Dataset information
    7. Unit of Analysis
    8. Target Features
    9. Prediction plan
    Question and Answers:
    `{total_text}`

    Please mark main heading in the document with a single # and all sub headers with a double ##
'''

answer = chain.run(use_case_prompt)
print('++++++++++++++++++++++++++++++++++++++++++')
print(answer)

#############################################################################
# Use case Document generation
#############################################################################

# Create an instance of a word document
doc = docx.Document()
  
lines = answer.split('\n')
for line in lines:
    if line.startswith('#'):
        line = line.split("#")[-1].strip()

        # Add a Title to the document 
        doc.add_heading(line, 0)
    elif line.startswith('##'):
        line = line.split("##")[-1].strip()
        
        doc.add_heading(f"{line}:", 0)
    else:
        # Adding paragraph
        doc.add_paragraph(line)
  
# Now save the document to a location 
doc.save('use_case_document.docx')
#############################################################################

#############################################################################
# Dataset generation
#############################################################################
# Extract relevant sections for the dataset
dataset_info_start = answer.find("## Dataset Information")
unit_of_analysis_start = answer.find("## Unit of Analysis")
target_feature_start = answer.find("## Target Feature")

dataset_info = answer[dataset_info_start + len("## Dataset Information"):unit_of_analysis_start]
unit_of_analysis = answer[unit_of_analysis_start + len("## Unit of Analysis"):target_feature_start]
target_feature = answer[target_feature_start + len("## Target Feature"):]

# Create a dataset dictionary
dataset = {
    "dataset_info": dataset_info,
    "unit_of_analysis": unit_of_analysis,
    "target_feature": target_feature
}

# Print the dataset dictionary
print("DATASET: ")
print(dataset)

# Create a DataFrame from the dataset dictionary
df = pd.DataFrame.from_dict(dataset, orient="index", columns=["Content"])

# Display the DataFrame
print(df)
#############################################################################

