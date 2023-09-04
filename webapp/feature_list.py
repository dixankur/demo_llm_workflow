import streamlit as st
from webapp.app_config import projects, project_stages, project_feature_list
from webapp.utils import provide_rating, chat_window, provide_appoval, update_status_and_audit

import pandas as pd
import json
import os
import time
import docx
from webapp.utils import save_json, add_sidebar

# show title
def render():
    prj_id = st.session_state['project_id']

    if prj_id in projects:
        prj_name = projects[prj_id]
    else:
        prj_name = 'Unknown Project'

    # sidebar
    add_sidebar(header='CompeersAI', prj_name=prj_name, 
            notes=[
                "Add a table or dataset.",
                "Select features from table or datset.",
                "Define how to combine tables or datasets."
            ])
    
    st.header(':red[Feature list selection]')
    st.subheader(':blue[Table details]', divider='rainbow')
    
    def prj_click_button():
        st.session_state['current_page'] = 'prj_home'

    if prj_id in project_feature_list:
        f_list = project_feature_list[prj_id]
        for t_details in f_list['tables']:
            st.subheader(t_details['name'])
            for col_val in t_details['columns']:
                st.checkbox(col_val[0], value=col_val[1], key=t_details['name']+'_'+col_val[0])
    
    st.divider()
    st.subheader(':blue[Add Table]', divider='rainbow')
    # Collect user input for table name and columns using st.form
    with st.form("user_input_form"):
        table_name = st.text_input(label="Enter the table name:")
        column_names = st.text_input(label="Enter column names (comma-separated):")
        submit_button = st.form_submit_button(label="Create Table")

    # Convert column names to a list
    column_list = [col.strip() for col in column_names.split(',')]

    # Display the table if the form is submitted and there's valid input
    if submit_button and table_name and column_list:
        table_names = []
        for table in project_feature_list['P001']['tables']:
            table_names.append(table.get('name'))
        if table_name in table_names:
            st.markdown(f":red[Table '{table_name}' already exists!!]")
        else:
            new_table = {}            
            ff_list = []
            # Display column names
            for col in column_list:
                ff_list.append([col, True])

            new_table['name'] = table_name
            new_table['path'] = f"{table_name}.csv"
            new_table['format'] = 'csv'
            new_table['columns'] = ff_list
            project_feature_list['P001']['tables'].append(new_table)

            # save json 
            save_json(project_feature_list, './datamodel/project_feature_list.json')

    st.divider()    
    st.subheader(':blue[Document Section]', divider='rainbow')
    sd1, sd2, sd3 = st.columns(3)
    with sd1:
        but1 = sd1.button('Generate Feature List document', type='secondary', use_container_width=True)
    if but1:
        doc = docx.Document()

        doc.add_heading("Feature List Document", 0)
        # read the list of tables
        with open ('./datamodel/project_feature_list.json', "r") as f:
            table_data = json.loads(f.read())
        # Extract table information
        for tab_details in table_data[prj_id]['tables']:
            doc.add_heading(f"Table: {tab_details['name']}", 0)
            doc.add_paragraph(f"Path: ./data/{prj_id}/{tab_details['path']}")
            doc.add_paragraph(f"Format: {tab_details['format']}")
            doc.add_paragraph(f"**Table columns:**")
            for col_idx, col_nm in enumerate(tab_details['columns']):
                if col_idx == 0:
                    doc.add_paragraph(col_nm[0])
                else:
                    doc.add_paragraph(col_nm[0])

        doc.save(f'./data/{prj_id}/feature_list_document.docx')
        update_status_and_audit(prj_id, "S02", "Inprogress", "Feature List Document generated")

    feature_list_doc_exists = os.path.exists(f"./data/{prj_id}/feature_list_document.docx")
    if feature_list_doc_exists:
        creation_time = os.path.getctime(f"./data/{prj_id}/feature_list_document.docx")
        # Convert the timestamps to datetime objects
        creation_datetime = time.ctime(creation_time)
        st.markdown(f":green[Feature List Document Created!! {creation_datetime}]")
        
    with sd2:
        if feature_list_doc_exists:
            with open(f"./data/{prj_id}/feature_list_document.docx", "rb") as docx_file:
                DOCXbyte = docx_file.read()
            if DOCXbyte:
                if st.download_button(label="Download Feature List Document", key='5',
                        data=DOCXbyte,
                        file_name=f"./data/{prj_id}/feature_list_document.docx",
                        mime='application/octet-stream', use_container_width=True):
                    st.markdown(f":green[Feature List Document Downloaded!!]")
        else:
            st.markdown(f":red[Feature List Document not generated yet!!]")

    with sd3:
        button_click = st.button('Upload Feature List Document', type='secondary', use_container_width=True)
    if button_click:
        uploaded_file = st.file_uploader("Choose Feature list document to upload", type=['docx'])
        if uploaded_file is not None:
            # To read file as bytes:
            #bytes_data = uploaded_file.getvalue()
            with open(f"./data/{prj_id}/feature_list_document.docx", "wb") as f: 
                f.write(uploaded_file.getbuffer())
            st.success("Saved File")

    if feature_list_doc_exists:
        provide_rating(prj_id, 'feature_list_doc')

    st.divider()
    st.markdown("<h2 style='color:tomato;font-size:150%;font-family:verdana'>Approve Feature List Document</h2>",
             unsafe_allow_html=True)
    apv_time = provide_appoval(prj_id, 'feature_list_doc_approval', 'Approve Feature List document')
    if apv_time:
        update_status_and_audit(prj_id, "S02", "Complete", "Feature List Document approved")
    st.divider()
    
    st.button('Back to Project Home', on_click=prj_click_button)

    #chat_window(prj_id, 'S02')
